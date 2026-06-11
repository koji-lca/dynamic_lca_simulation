"""
build_word_report.py — HTMLレポート → Word (.docx) 変換スクリプト

使用方法:
    python build_word_report.py
    （または .venv を有効化して実行）

出力: ../SuMPO動的LCA調査報告書.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
import lxml.html

# ---------------------------------------------------------------------------
# 定数
# ---------------------------------------------------------------------------
HERE = Path(__file__).parent
OUTPUT = HERE.parent / "SuMPO動的LCA調査報告書.docx"

C_NAVY  = RGBColor(0x19, 0x39, 0x50)
C_DBLUE = RGBColor(0x12, 0x54, 0x86)
C_BLUE  = RGBColor(0x00, 0x57, 0xBA)
C_GRAY  = RGBColor(0x55, 0x55, 0x55)

PAGE_FILES = [
    ("調査ブロック①　政策及び制度調査", [
        HERE / "survey1" / "ch01.html",
        HERE / "survey1" / "ch02.html",
        HERE / "survey1" / "ch03.html",
        HERE / "survey1" / "ch04.html",
    ]),
    ("調査ブロック②　技術及び手法調査", [
        HERE / "survey2" / "tech01.html",
        HERE / "survey2" / "tech02.html",
        HERE / "survey2" / "tech03.html",
    ]),
    ("PMO評価", [
        HERE / "pmo" / "evaluation_survey1.html",
        HERE / "pmo" / "evaluation_survey2.html",
    ]),
]


# ---------------------------------------------------------------------------
# XML ユーティリティ
# ---------------------------------------------------------------------------

def set_east_asian_font(element, font_name: str):
    """スタイルまたは Run に東アジアフォントを設定する。"""
    if hasattr(element, '_element'):
        rPr = element._element.get_or_add_rPr()
    else:
        rPr = element._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:eastAsiaTheme'), 'minorEastAsia')


def set_cell_shading(cell, fill_hex: str, font_color_hex: str = None):
    """テーブルセルの背景色を設定する。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.upper())
    tcPr.append(shd)
    if font_color_hex:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(font_color_hex)


def add_bottom_border_to_para_style(style):
    """段落スタイルに下線ボーダーを追加する（Heading 1 用）。"""
    pPr = style._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '0057BA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_left_border_to_para_style(style):
    """段落スタイルに左ボーダーを追加する（Heading 2 用）。"""
    pPr = style._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '0057BA')
    pBdr.append(left)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# スタイル設定
# ---------------------------------------------------------------------------

def setup_styles(doc: Document):
    """Wordドキュメントのスタイルを設定する。"""
    styles = doc.styles

    # Normal スタイル
    normal = styles['Normal']
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    set_east_asian_font(normal, 'MS Mincho')

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = C_NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    set_east_asian_font(h1, 'MS Gothic')
    try:
        add_bottom_border_to_para_style(h1)
    except Exception:
        pass

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = C_DBLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    set_east_asian_font(h2, 'MS Gothic')
    try:
        add_left_border_to_para_style(h2)
    except Exception:
        pass

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = C_NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    set_east_asian_font(h3, 'MS Gothic')

    # Heading 4
    h4 = styles['Heading 4']
    h4.font.size = Pt(10.5)
    h4.font.bold = True
    h4.font.color.rgb = C_DBLUE
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.keep_with_next = True
    set_east_asian_font(h4, 'MS Gothic')

    # Intense Quote（blockquote 用）
    try:
        iq = styles['Intense Quote']
        iq.font.italic = False
        iq.font.color.rgb = C_DBLUE
        iq.paragraph_format.left_indent = Cm(1)
    except Exception:
        pass

    # Caption スタイル
    try:
        cap = styles['Caption']
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = C_GRAY
    except Exception:
        pass


# ---------------------------------------------------------------------------
# 表紙ページ
# ---------------------------------------------------------------------------

def add_cover_page(doc: Document):
    """表紙を追加する。"""
    # 余白を確保するための空行
    for _ in range(8):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SuMPO 動的LCA調査報告書")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = C_NAVY
    set_east_asian_font(run, 'MS Gothic')

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("フランスRE2020における動的LCA手法の調査・分析")
    run2.font.size = Pt(14)
    run2.font.color.rgb = C_DBLUE
    set_east_asian_font(run2, 'MS Mincho')

    for _ in range(4):
        doc.add_paragraph()

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = org.add_run("一般社団法人サステナブル経営推進機構（SuMPO）")
    run3.font.size = Pt(12)
    set_east_asian_font(run3, 'MS Mincho')

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = date.add_run("2026年6月")
    run4.font.size = Pt(12)
    set_east_asian_font(run4, 'MS Mincho')

    doc.add_page_break()


# ---------------------------------------------------------------------------
# インラインコンテンツ処理
# ---------------------------------------------------------------------------

def _split_inline_math(text: str):
    """テキストをインライン数式($...$)で分割してリストを返す。
    偶数インデックス: 通常テキスト、奇数インデックス: 数式
    """
    return re.split(r'\$(.+?)\$', text, flags=re.DOTALL)


def _add_text_runs(para, text: str, bold=False, italic=False, mono=False):
    """テキストを通常/数式に分割してRunとして追加する。"""
    parts = _split_inline_math(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            # インライン数式: 等幅フォントで表示
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            run.bold = bold
            run.italic = italic
            if mono:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            else:
                set_east_asian_font(run, 'MS Mincho')


def add_runs_from_element(para, el, bold=False, italic=False):
    """lxml要素からインラインコンテンツを処理してRunを追加する。"""
    # el.text（開始タグ直後のテキスト）
    if el.text:
        _add_text_runs(para, el.text, bold=bold, italic=italic)

    for child in el:
        # Comment ノードはスキップ
        if callable(child.tag):
            if child.tail:
                _add_text_runs(para, child.tail, bold=bold, italic=italic)
            continue

        tag = child.tag.lower() if isinstance(child.tag, str) else ''

        if tag in ('strong', 'b'):
            add_runs_from_element(para, child, bold=True, italic=italic)
        elif tag in ('em', 'i'):
            add_runs_from_element(para, child, bold=bold, italic=True)
        elif tag == 'code':
            text = child.text_content()
            if text:
                run = para.add_run(text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        elif tag == 'br':
            para.add_run('\n')
        elif tag == 'a':
            add_runs_from_element(para, child, bold=bold, italic=italic)
        elif tag == 'span':
            add_runs_from_element(para, child, bold=bold, italic=italic)
        elif tag in ('sub', 'sup'):
            add_runs_from_element(para, child, bold=bold, italic=italic)
        else:
            # その他の要素はテキスト内容のみ
            text = child.text_content()
            if text:
                _add_text_runs(para, text, bold=bold, italic=italic)

        # 子要素の後ろのテキスト（tail）
        if child.tail:
            _add_text_runs(para, child.tail, bold=bold, italic=italic)


# ---------------------------------------------------------------------------
# 要素コンバーター
# ---------------------------------------------------------------------------

def convert_heading(doc: Document, el):
    """h1〜h4 → Word Heading スタイル。"""
    level = int(el.tag[1])
    text = el.text_content().strip()
    if not text:
        return
    para = doc.add_heading(text, level=min(level, 4))
    # 見出しに東アジアフォントを設定
    for run in para.runs:
        set_east_asian_font(run, 'MS Gothic')


def convert_paragraph(doc: Document, el, style=None):
    """p → Word 段落。"""
    # 空の段落はスキップ
    if not el.text_content().strip():
        return
    para = doc.add_paragraph(style=style)
    add_runs_from_element(para, el)


def convert_list(doc: Document, el, ordered=False):
    """ul/ol → リスト段落。"""
    style = 'List Number' if ordered else 'List Bullet'
    for li in el.findall('.//li'):
        if li.getparent() != el:
            continue  # ネストされたリストは親のみ処理
        para = doc.add_paragraph(style=style)
        add_runs_from_element(para, li)


def convert_blockquote(doc: Document, el):
    """blockquote → Intense Quote スタイル。"""
    for child in el:
        if callable(child.tag):
            continue
        tag = child.tag.lower()
        if tag == 'p':
            try:
                para = doc.add_paragraph(style='Intense Quote')
            except Exception:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
            add_runs_from_element(para, child)
        elif tag in ('ul', 'ol'):
            convert_list(doc, child, ordered=(tag == 'ol'))
        elif tag == 'blockquote':
            convert_blockquote(doc, child)


def _get_table_rows_and_max_cols(el):
    """テーブルの全行と最大列数を返す。"""
    rows = el.findall('.//tr')
    if not rows:
        return rows, 0
    max_cols = 0
    for tr in rows:
        cells = tr.findall('th') + tr.findall('td')
        col_count = sum(int(c.get('colspan', 1)) for c in cells)
        max_cols = max(max_cols, col_count)
    return rows, max_cols


def _fill_cell(word_cell, cell_el):
    """テーブルセルにコンテンツを追加する。"""
    para = word_cell.paragraphs[0]
    # セル内に <p> がある場合はそれぞれ処理
    p_children = [c for c in cell_el if not callable(c.tag) and c.tag.lower() == 'p']
    if p_children:
        add_runs_from_element(para, p_children[0])
        for extra_p in p_children[1:]:
            new_para = word_cell.add_paragraph()
            add_runs_from_element(new_para, extra_p)
    else:
        add_runs_from_element(para, cell_el)


def convert_table(doc: Document, el):
    """table → Word テーブル。"""
    rows, max_cols = _get_table_rows_and_max_cols(el)
    if not rows or max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for row_idx, tr in enumerate(rows):
        # ヘッダー行の判定
        parent_tag = tr.getparent().tag.lower() if tr.getparent() is not None else ''
        cells_el = [c for c in tr if not callable(c.tag) and c.tag.lower() in ('th', 'td')]
        is_header = (parent_tag == 'thead') or all(c.tag.lower() == 'th' for c in cells_el)

        col_idx = 0
        for cell_el in cells_el:
            if col_idx >= max_cols:
                break
            word_cell = table.cell(row_idx, col_idx)
            _fill_cell(word_cell, cell_el)

            # colspan のマージ処理
            colspan = int(cell_el.get('colspan', 1))
            if colspan > 1:
                end_col = min(col_idx + colspan - 1, max_cols - 1)
                if end_col > col_idx:
                    word_cell.merge(table.cell(row_idx, end_col))

            if is_header:
                set_cell_shading(word_cell, '193950', 'FFFFFF')
                for para in word_cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            elif row_idx % 2 == 1:
                set_cell_shading(word_cell, 'F4F6F8')

            col_idx += colspan

    # テーブル後の空行
    doc.add_paragraph()


def convert_math_block(doc: Document, el):
    """div.math-block → 等幅フォントの数式ブロック。"""
    raw = el.text_content().strip()
    # $$ デリミタを除去
    latex = re.sub(r'^\$\$\s*', '', raw)
    latex = re.sub(r'\s*\$\$$', '', latex)
    latex = latex.strip()
    if not latex:
        return

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(latex)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def convert_figure(doc: Document, html_path: Path, img_src: str, caption_text: str):
    """画像とキャプションを追加する。"""
    # 相対パスを絶対パスに変換
    abs_img_path = (html_path.parent / img_src).resolve()

    if abs_img_path.exists():
        try:
            doc.add_picture(str(abs_img_path), width=Cm(14))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"  [警告] 画像読み込みエラー: {abs_img_path} — {e}", file=sys.stderr)
    else:
        print(f"  [警告] 画像ファイルが見つかりません: {abs_img_path}", file=sys.stderr)

    if caption_text.strip():
        try:
            cap_para = doc.add_paragraph(style='Caption')
        except Exception:
            cap_para = doc.add_paragraph()
            for run in cap_para.runs:
                run.font.size = Pt(9)
                run.font.italic = True
        cap_run = cap_para.add_run(caption_text.strip())
        set_east_asian_font(cap_run, 'MS Mincho')
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def convert_codehilite(doc: Document, el):
    """div.codehilite → Courier New 等幅ブロック。"""
    pre = el.find('.//pre')
    if pre is None:
        return
    code_text = pre.text_content()
    if not code_text.strip():
        return

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)


def convert_comment_fig(doc: Document, html_path: Path, comment_text: str):
    """<!-- FIG:filename|caption --> コメントを図として処理する。"""
    match = re.match(r'\s*FIG:([^|]+)\|(.+)', comment_text)
    if not match:
        return
    img_filename = match.group(1).strip()
    caption = match.group(2).strip()
    img_src = f"../assets/charts/{img_filename}"
    convert_figure(doc, html_path, img_src, caption)


# ---------------------------------------------------------------------------
# ページ処理
# ---------------------------------------------------------------------------

def _get_classes(el) -> list:
    """要素の class 属性をリストとして返す。"""
    return (el.get('class') or '').split()


def process_element(doc: Document, html_path: Path, el):
    """単一の lxml 要素を処理してWordドキュメントに追加する。"""
    # コメントノード
    if callable(el.tag):
        comment_text = el.text or ''
        if 'FIG:' in comment_text:
            convert_comment_fig(doc, html_path, comment_text)
        return

    tag = el.tag.lower()
    classes = _get_classes(el)

    # スキップ: ナビゲーション要素
    if 'breadcrumb' in classes or 'section-badge' in classes:
        return
    if tag == 'hr':
        return
    if tag == 'script':
        return

    # 見出し
    if tag in ('h1', 'h2', 'h3', 'h4', 'h5'):
        convert_heading(doc, el)

    # 段落
    elif tag == 'p':
        # p の直下に div.math-block がある場合は math-block として処理
        # (lxmlはこれをhoist済みなので通常のpとして来る)
        convert_paragraph(doc, el)

    # リスト
    elif tag == 'ul':
        convert_list(doc, el, ordered=False)
    elif tag == 'ol':
        convert_list(doc, el, ordered=True)

    # Blockquote
    elif tag == 'blockquote':
        convert_blockquote(doc, el)

    # テーブル
    elif tag == 'table':
        convert_table(doc, el)

    # 数式ブロック
    elif tag == 'div' and 'math-block' in classes:
        convert_math_block(doc, el)

    # 図ブロック
    elif tag == 'div' and 'fig-block' in classes:
        img = el.find('.//img')
        if img is not None:
            img_src = img.get('src', '')
            cap_el = el.find('./p[@class="fig-caption"]')
            if cap_el is None:
                cap_el = el.find('.//p')
            caption = cap_el.text_content() if cap_el is not None else ''
            convert_figure(doc, html_path, img_src, caption)

    # コードブロック
    elif tag == 'div' and 'codehilite' in classes:
        convert_codehilite(doc, el)

    # その他の div: 子要素を再帰処理
    elif tag == 'div':
        for child in el:
            process_element(doc, html_path, child)


def process_page(doc: Document, html_path: Path):
    """1つのHTMLファイルを処理してWordドキュメントに追加する。"""
    print(f"  処理中: {html_path.name}", file=sys.stderr)

    content = html_path.read_text(encoding='utf-8')
    root = lxml.html.fromstring(content)

    # メインコンテンツ div#content を取得
    content_div = root.find('.//*[@id="content"]')
    if content_div is None:
        print(f"  [警告] div#content が見つかりません: {html_path.name}", file=sys.stderr)
        return

    for el in content_div:
        process_element(doc, html_path, el)


# ---------------------------------------------------------------------------
# メイン
# ---------------------------------------------------------------------------

def main():
    print("Word レポート生成を開始します...", file=sys.stderr)

    doc = Document()

    # A4 ページ設定
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    setup_styles(doc)
    add_cover_page(doc)

    first_page = True
    for section_label, html_paths in PAGE_FILES:
        # セクション見出し（調査ブロック区切り）
        if not first_page:
            doc.add_page_break()
        first_page = False

        sec_para = doc.add_paragraph()
        sec_run = sec_para.add_run(section_label)
        sec_run.font.size = Pt(18)
        sec_run.font.bold = True
        sec_run.font.color.rgb = C_NAVY
        set_east_asian_font(sec_run, 'MS Gothic')
        sec_para.paragraph_format.space_after = Pt(12)

        # 各HTMLページ
        for i, html_path in enumerate(html_paths):
            if not html_path.exists():
                print(f"  [警告] ファイルが見つかりません: {html_path}", file=sys.stderr)
                continue
            if i > 0:
                doc.add_page_break()
            process_page(doc, html_path)

    doc.save(str(OUTPUT))
    print(f"\n完了: {OUTPUT}", file=sys.stderr)


if __name__ == '__main__':
    main()
